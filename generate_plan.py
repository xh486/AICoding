"""
生成《AI应用开发学习计划书》Word文档
输出到桌面
"""
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

doc = Document()

# ── 全局样式设置 ──
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 页边距
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)


def add_heading(text, level=1):
    """添加带字体设置的标题"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return h


def add_para(text, bold=False, indent=False, font_size=None):
    """添加段落"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if bold:
        run.bold = True
    if font_size:
        run.font.size = Pt(font_size)
    return p


def add_bullet(text, level=0):
    """添加项目符号"""
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(11)
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * level)
    return p


def add_tip(text):
    """添加提示/重点段落"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run('💡 ' + text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x33, 0x66, 0x99)
    return p


# ═══════════════════════════════════════════
# 封面
# ═══════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('AI应用开发学习计划书')
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('从零到就业 · 四个月路线图')
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('适用对象：大四在校生 / 应届毕业生\n目标岗位：AI应用开发工程师\n学习周期：4个月（16周）\n编写日期：2026年6月')
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
run.font.size = Pt(11)

doc.add_page_break()

# ═══════════════════════════════════════════
# 目录页（手动）
# ═══════════════════════════════════════════
add_heading('目录', level=1)
toc_items = [
    '一、行业认知 — AI应用开发是什么',
    '二、技术栈清单 — 学什么，不学什么',
    '三、四个月学习路线图',
    '四、每周详细计划（第1-16周）',
    '五、项目里程碑',
    '六、开发环境搭建指南',
    '七、资源推荐',
    '八、求职策略',
]
for item in toc_items:
    add_para(item, font_size=12)
doc.add_page_break()

# ═══════════════════════════════════════════
# 一、行业认知
# ═══════════════════════════════════════════
add_heading('一、行业认知 — AI应用开发是什么', level=1)

add_heading('1.1 这个行业到底在做什么', level=2)
add_para('AI应用开发 ≠ 训练模型。你不是研究者，你是工程师。你的工作是：把大语言模型（LLM）的能力封装成可用的产品。')
add_para('')
add_para('日常核心工作就三件事：')
add_bullet('调用LLM API（OpenAI、Claude、通义千问等），设计Prompt控制输出质量')
add_bullet('搭建RAG（检索增强生成）系统，让AI能回答基于企业私有文档的问题')
add_bullet('开发AI Agent，让模型能调用工具、执行多步任务')
add_para('')
add_para('说白了：你不是造发动机的人，你是用发动机造汽车的人。')

add_heading('1.2 为什么现在是最好的入场时机', level=2)
add_bullet('门槛急速降低：2024年还需要会PyTorch、会微调；2026年用Vercel AI SDK + 一个API Key就能搭起AI应用')
add_bullet('需求爆炸、供给严重不足：几乎每家公司都在想办法"接入AI"，但真正能落地的工程师太少')
add_bullet('薪资溢价明显：同等经验的普通前端/后端 vs AI应用开发，后者有明显溢价')
add_bullet('技术栈收敛：行业逐渐形成以TypeScript/Python为主的标准技术栈，学了就能用')

add_heading('1.3 市场到底要什么样的人', level=2)
add_bullet('能独立把一个AI产品从零搭到上线的人，不是"懂原理"的人')
add_bullet('工程能力 > 算法知识。你会写干净的代码、能处理边界情况、能部署上线，比你会推公式重要10倍')
add_bullet('GitHub上有可运行的项目 > 简历上写"熟悉大模型原理"')

add_tip('记住这句话：这个行业缺的不是懂AI的人，是能把AI变成产品的人。')

doc.add_page_break()

# ═══════════════════════════════════════════
# 二、技术栈清单
# ═══════════════════════════════════════════
add_heading('二、技术栈清单 — 学什么，不学什么', level=1)

add_heading('2.1 必须掌握的（按优先级排序）', level=2)

# 表格
table = doc.add_table(rows=7, cols=4)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['优先级', '技术', '学多久', '为什么重要']
rows_data = [
    ['P0（最高）', 'TypeScript + Next.js', '4-6周', 'AI应用90%是Web产品，Next.js是全栈框架的事实标准。Vercel AI SDK、LangChain.js都在这上面'],
    ['P0（最高）', 'Python', '2周', '数据处理、LangChain、后端Agent的母语。不需要精通，能写API就行'],
    ['P1', 'LLM 基础概念', '1周', 'Token / Temperature / System Prompt / Few-shot / RAG / Embedding / Function Calling — 这是AI开发的"专业课"'],
    ['P1', 'Prompt Engineering', '持续学', '这是AI时代的"新编程语言"。Prompt写不好，再好的模型也白搭'],
    ['P1', 'Vector DB + RAG', '1-2周', '90%企业级AI应用的核心：把企业文档变成可问答的知识库。主攻Supabase pgvector或Pinecone'],
    ['P2', 'PostgreSQL + SQL', '2-3周（如果不会）', 'AI应用的持久化层，绕不开。会写查询、理解索引就够了'],
]

for i, h_text in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h_text
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for r, row_data in enumerate(rows_data):
    for c, cell_text in enumerate(row_data):
        cell = table.rows[r + 1].cells[c]
        cell.text = cell_text
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

add_heading('2.2 加分项（让你在竞争者中脱颖而出的）', level=2)
add_bullet('Docker + 部署（Vercel / Railway / 阿里云）：能把自己写的东西跑在服务器上——这是学生最缺的能力')
add_bullet('Git 协作：不只是 add/commit/push，要会 branch、rebase、code review、PR 流程')
add_bullet('LangChain 或 LlamaIndex：市场占有率最高的两个AI编排框架，至少深入了解一个')
add_bullet('Agent 开发：2026年最热的词。理解 tool use、multi-agent orchestration、plan-and-execute 模式')
add_bullet('FastAPI（Python）或 Hono（TypeScript）：任何一个后端框架')

add_heading('2.3 现在不要碰的（省时间）', level=2)
add_para('❌ 从零学机器学习/深度学习理论 — 不是没用，但你目标是就业不是读博', bold=True)
add_para('❌ CUDA / GPU编程 — 你是用模型的人，不是训模型的人', bold=True)
add_para('❌ "学完一门语言再学下一门" — TypeScript和Python并行学，不需要精通一门才开始另一门', bold=True)
add_para('❌ 追逐每一个新模型/新框架 — 跟一个主流技术栈深耕，别当工具收集者', bold=True)

add_heading('2.4 2026年推荐技术栈组合', level=2)
add_bullet('前端/全栈：Next.js (App Router) + TypeScript + Vercel AI SDK + Tailwind CSS')
add_bullet('后端：FastAPI (Python) 或 Next.js API Routes')
add_bullet('LLM：Claude API (Anthropic) + OpenAI API')
add_bullet('向量数据库：Supabase (pgvector) — 免费额度够学习用')
add_bullet('部署：Vercel（前端）+ Railway 或 阿里云ECS（后端）')
add_bullet('数据库：Supabase PostgreSQL 或 PlanetScale')

doc.add_page_break()

# ═══════════════════════════════════════════
# 三、四个月学习路线图
# ═══════════════════════════════════════════
add_heading('三、四个月学习路线图', level=1)

months = [
    {
        'title': '第一个月：基础篇 — Web开发 + AI初体验',
        'goal': '能用Next.js搭网站，能调大模型API做出一个AI聊天页面',
        'tasks': [
            'Next.js基础：App Router、Server Component、路由、布局',
            'Tailwind CSS：写出能看的界面',
            'API Routes：写后端接口',
            '接入Claude API或OpenAI API',
            '实现一个带对话历史的AI聊天页面',
            '用环境变量管理API Key',
            '部署到Vercel，获得一个公网URL',
        ],
        'deliverable': '✅ 一个部署在公网上的AI聊天应用，别人能打开用',
    },
    {
        'title': '第二个月：RAG篇 — 让AI读懂你的文档',
        'goal': '做一个"上传PDF然后能问答"的完整应用',
        'tasks': [
            '理解Embedding是什么（文本→向量→语义搜索）',
            '理解Chunking策略（怎么把长文档切块）',
            '用Supabase搭建pgvector向量数据库',
            '实现PDF上传→解析→Embedding→存入向量库的完整流水线',
            '实现"提问→检索相关段落→拼接上下文→LLM生成答案"的RAG流程',
            '优化检索质量：调整chunk大小、尝试不同的相似度算法',
        ],
        'deliverable': '✅ PDF知识库问答系统，能上传任意PDF并基于内容问答',
    },
    {
        'title': '第三个月：Agent篇 — 让AI能动手做事',
        'goal': '做一个能调用工具、执行多步任务的Agent',
        'tasks': [
            '理解Function Calling / Tool Use机制',
            '让Agent能调用外部API（天气查询、网页搜索等）',
            '实现多步推理Agent Loop（思考→行动→观察→再思考）',
            '尝试LangChain或Vercel AI SDK的Agent模块',
            '做一个"AI助手"：能查天气、搜资料、算数学题',
            '理解Agent的安全边界（权限控制、输入校验）',
        ],
        'deliverable': '✅ 一个能使用多种工具的AI Agent，能自主完成多步任务',
    },
    {
        'title': '第四个月：成品篇 — 拿得出手的项目 + 求职冲刺',
        'goal': '完成一个完整作品，投出第一批简历',
        'tasks': [
            '选定一个方向做最终作品（记账助手/英语陪练/简历优化/代码Review工具）',
            '完整实现：用户系统、数据持久化、AI核心功能、前端界面',
            '写README、录Demo视频',
            '整理GitHub Profile，至少3个可运行的项目',
            '写简历，突出"独立交付AI产品"的能力',
            '开始投简历、参加面试',
        ],
        'deliverable': '✅ 一个完整的产品级AI应用 + 优化的简历 + 开始面试',
    },
]

for i, month in enumerate(months):
    add_heading(month['title'], level=2)
    add_para(f'🎯 目标：{month["goal"]}', bold=True)
    add_para('具体任务：')
    for task in month['tasks']:
        add_bullet(task)
    add_para('')
    add_para(month['deliverable'], bold=True)
    if i < len(months) - 1:
        add_para('')

doc.add_page_break()

# ═══════════════════════════════════════════
# 四、每周详细计划
# ═══════════════════════════════════════════
add_heading('四、每周详细计划（第1-16周）', level=1)
add_para('说明：按每天2-3小时有效学习时间估算。如果全天学习，速度可以翻倍。', font_size=10)

weeks = [
    # 第1个月
    {'week': '第1周', 'theme': 'TypeScript 快速上手', 'tasks': [
        'Day1-2：安装Node.js、VS Code，配置开发环境',
        'Day3-4：TypeScript基础语法 — 类型注解、interface、泛型（够用就行，不用精通）',
        'Day5-6：异步编程 — async/await、Promise、fetch',
        'Day7：写一个小练习：用fetch请求一个公开API并打印结果',
    ], 'deliverable': '能看懂TypeScript代码，能用fetch调API'},
    {'week': '第2周', 'theme': 'Next.js入门 + 第一个页面', 'tasks': [
        'Day1-2：npx create-next-app，理解项目结构',
        'Day3-4：App Router路由系统 — page.tsx、layout.tsx、动态路由',
        'Day5-6：Server Component vs Client Component，理解RSC',
        'Day7：用Tailwind CSS写一个个人主页',
    ], 'deliverable': '一个部署到Vercel的个人主页'},
    {'week': '第3周', 'theme': 'API Routes + 接入AI', 'tasks': [
        'Day1-2：Next.js API Routes — 写第一个POST接口',
        'Day3-4：注册Claude API / OpenAI API账号，获取API Key',
        'Day5-6：用环境变量管理密钥，在API Route中调用LLM',
        'Day7：前端页面发送请求到API Route，展示AI回复',
    ], 'deliverable': '前后端联通的AI对话接口'},
    {'week': '第4周', 'theme': 'AI聊天应用 + 上线', 'tasks': [
        'Day1-3：实现多轮对话历史管理',
        'Day4-5：优化UI — 消息气泡、加载状态、错误处理',
        'Day6：部署到Vercel，配置环境变量',
        'Day7：测试、修复部署问题、发给朋友试用',
    ], 'deliverable': '公网可访问的AI聊天应用'},

    # 第2个月
    {'week': '第5周', 'theme': 'Embedding 与向量数据库', 'tasks': [
        'Day1-2：理解Embedding — 文字怎么变成向量、余弦相似度',
        'Day3-4：注册Supabase，开启pgvector扩展',
        'Day5-6：用OpenAI/Claude Embedding API生成向量，存入Supabase',
        'Day7：实现"输入问题→生成问题向量→检索最相似的文档片段"',
    ], 'deliverable': '能跑通的向量检索demo'},
    {'week': '第6周', 'theme': 'PDF处理 + Chunking', 'tasks': [
        'Day1-2：用pdf-parse或PyPDF2解析PDF文本',
        'Day3-4：实现Chunking策略 — 固定大小切块、语义切块',
        'Day5-6：搭建"上传PDF→解析→切块→生成Embedding→存储"流水线',
        'Day7：完整跑通流程，验证检索效果',
    ], 'deliverable': 'PDF上传到向量检索的完整流水线'},
    {'week': '第7周', 'theme': 'RAG问答系统', 'tasks': [
        'Day1-2：实现RAG核心逻辑：检索→拼接Context→构造Prompt→LLM生成',
        'Day3-4：处理引用来源 — 让AI回答时标注来自PDF的哪一页',
        'Day5-6：优化体验 — 流式输出、对话历史、Source高亮',
        'Day7：部署上线，用真实文档测试',
    ], 'deliverable': '能用的PDF知识库问答系统'},
    {'week': '第8周', 'theme': 'RAG优化 + 巩固', 'tasks': [
        'Day1-2：实验不同chunk大小对检索质量的影响',
        'Day3-4：加入Hyde（假设文档嵌入）、重排序等优化手段',
        'Day5-6：处理Edge Case — 空文档、超长文档、非中文文档',
        'Day7：写一篇RAG系统的技术文档，整理到GitHub README',
    ], 'deliverable': '优化后的RAG系统 + 技术文档'},

    # 第3个月
    {'week': '第9周', 'theme': 'Function Calling基础', 'tasks': [
        'Day1-2：理解Function Calling原理 — Tool定义、模型如何选择调用',
        'Day3-4：定义第一个Tool — 查天气（调用公开天气API）',
        'Day5-6：实现完整的Tool Call循环 — LLM决定调用→执行函数→返回结果→LLM总结',
        'Day7：让AI能根据用户问题自动决定要不要查天气',
    ], 'deliverable': '一个能调用天气查询工具的Agent'},
    {'week': '第10周', 'theme': '多工具Agent', 'tasks': [
        'Day1-2：添加第2个Tool — 网页搜索',
        'Day3-4：添加第3个Tool — 计算器/数学运算',
        'Day5-6：处理多工具调用场景 — Agent自主选择用哪个工具',
        'Day7：处理工具调用失败的容错逻辑',
    ], 'deliverable': '能使用3种工具的Agent'},
    {'week': '第11周', 'theme': 'Agent Loop + 多步推理', 'tasks': [
        'Day1-2：理解ReAct模式 — Thought → Action → Observation → Thought',
        'Day3-4：实现Agent Loop — 循环直到完成任务或达到步数上限',
        'Day5-6：做一个复杂任务测试：如"查一下北京今天天气，如果是晴天就推荐户外活动"',
        'Day7：加入终止条件、超时控制、Token消耗监控',
    ], 'deliverable': '能自主完成多步推理的Agent'},
    {'week': '第12周', 'theme': 'Agent框架 + 项目整理', 'tasks': [
        'Day1-3：尝试Vercel AI SDK的useChat + toolInvocations或LangChain Agent',
        'Day4-5：理解Agent安全 — Prompt注入防护、工具权限最小化',
        'Day6-7：整理前三个月的所有项目到GitHub，确保README完善',
    ], 'deliverable': '整洁的GitHub主页，至少3个项目可运行'},

    # 第4个月
    {'week': '第13周', 'theme': '选定最终项目方向', 'tasks': [
        'Day1-2：头脑风暴，选1个有真实使用场景的方向：',
        '  • 简历优化助手（上传简历→AI分析→给出改进建议）',
        '  • AI英语陪练（语音/文字对话，纠正语法和发音）',
        '  • 个人记账AI助手（语音输入→自动分类→月度分析）',
        '  • 代码Review助手（提交代码→AI审查→给出建议）',
        'Day3-4：确定技术方案、画架构图、建GitHub仓库',
        'Day5-7：搭建项目骨架 — 数据库设计、页面路由、基础组件',
    ], 'deliverable': '最终项目的技术方案 + 项目骨架'},
    {'week': '第14周', 'theme': '核心功能开发', 'tasks': [
        'Day1-3：实现核心AI功能（如简历分析的Prompt链）',
        'Day4-5：实现用户系统（至少支持GitHub OAuth登录）',
        'Day6-7：数据持久化 — 用户数据、历史记录存入数据库',
    ], 'deliverable': '核心功能可运行'},
    {'week': '第15周', 'theme': '打磨 + 上线', 'tasks': [
        'Day1-2：UI/UX打磨 — 加载状态、空状态、错误提示、响应式',
        'Day3-4：性能优化 — 流式输出、缓存、图片优化',
        'Day5：完整部署上线，配置域名（如果有）',
        'Day6：写详细的README — 功能截图、技术栈、本地运行步骤',
        'Day7：录一段2分钟的Demo视频',
    ], 'deliverable': '产品上线 + README + Demo视频'},
    {'week': '第16周', 'theme': '求职冲刺', 'tasks': [
        'Day1-2：写简历 — 突出"独立交付了4个AI产品"，附上GitHub链接',
        'Day3-4：准备面试 — AI基础概念、项目技术细节、系统设计题',
        'Day5-7：开始投递 — 拉勾/Boss直聘/脉脉/LinkedIn，目标：AI应用开发/全栈工程师',
    ], 'deliverable': '简历投出 ≥ 20份，获得第一批面试机会'},
]

for w in weeks:
    add_heading(f'{w["week"]}：{w["theme"]}', level=2)
    for task in w['tasks']:
        add_bullet(task)
    add_para(f'📌 本周可交付成果：{w["deliverable"]}', bold=True)

doc.add_page_break()

# ═══════════════════════════════════════════
# 五、项目里程碑
# ═══════════════════════════════════════════
add_heading('五、项目里程碑', level=1)
add_para('四个关键里程碑，每个都是你简历上的一行项目经历：')

milestones = [
    ('🏆 里程碑1（第4周末）', 'AI聊天应用上线',
     '技术栈：Next.js + Claude API + Vercel\n功能：多轮对话、历史管理、流式输出\n证明你能：搭全栈应用、接大模型API、部署上线'),
    ('🏆 里程碑2（第8周末）', 'PDF知识库问答系统',
     '技术栈：Next.js + Supabase pgvector + Embedding API\n功能：PDF上传解析、语义检索、RAG问答、引用来源\n证明你能：做RAG、管向量数据库、处理非结构化数据'),
    ('🏆 里程碑3（第12周末）', '多工具AI Agent',
     '技术栈：Vercel AI SDK / LangChain + Function Calling\n功能：多工具调用、多步推理、自主决策\n证明你能：开发Agent、理解LLM推理机制、处理复杂交互'),
    ('🏆 里程碑4（第16周末）', '个人作品集项目',
     '技术栈：完整技术栈 + 用户系统 + 数据库\n功能：完整产品、用户认证、数据持久化、线上可访问\n证明你能：独立从零交付产品级应用'),
]

for title, name, desc in milestones:
    add_heading(title, level=2)
    add_para(name, bold=True, font_size=12)
    for line in desc.split('\n'):
        add_bullet(line)

doc.add_page_break()

# ═══════════════════════════════════════════
# 六、开发环境搭建指南
# ═══════════════════════════════════════════
add_heading('六、开发环境搭建指南', level=1)

add_heading('6.1 第一天：装好你的工具', level=2)
add_para('按顺序来，不要跳步：')

add_para('1. Node.js（18+版本）', bold=True)
add_bullet('去 nodejs.org 下载LTS版本，安装后终端输入 node -v 验证')

add_para('2. VS Code', bold=True)
add_bullet('安装以下扩展：ES7+ React/Redux/React-Native snippets、Tailwind CSS IntelliSense、Prettier、GitHub Copilot（或免费替代：Continue）')

add_para('3. Git + GitHub', bold=True)
add_bullet('注册GitHub账号，本地配置 git config --global user.name 和 user.email')

add_para('4. 终端', bold=True)
add_bullet('Windows用 Git Bash 或 Windows Terminal，不要用cmd')

add_para('5. Python（可选，第2个月用）', bold=True)
add_bullet('python.org下载3.11+版本，或者用Anaconda')

add_heading('6.2 前两周要注册的账号', level=2)
add_bullet('GitHub（代码托管）')
add_bullet('Vercel（前端部署，免费）')
add_bullet('Anthropic Console（Claude API）或 OpenAI Platform')
add_bullet('Supabase（数据库 + 向量存储，免费额度）')

add_heading('6.3 必会的终端命令', level=2)
add_bullet('npx create-next-app@latest my-app — 创建Next.js项目')
add_bullet('npm run dev — 启动开发服务器')
add_bullet('git add / commit / push — 提交代码到GitHub')
add_bullet('vercel deploy — 一键部署到Vercel')

doc.add_page_break()

# ═══════════════════════════════════════════
# 七、资源推荐
# ═══════════════════════════════════════════
add_heading('七、资源推荐', level=1)

add_heading('7.1 在线课程（免费为主）', level=2)
add_bullet('Next.js官方文档（nextjs.org/docs）— 最好的Next.js教程，没有之一')
add_bullet('Vercel AI SDK文档（sdk.vercel.ai）— AI应用开发的上手文档')
add_bullet('Anthropic Cookbook（github.com/anthropics/anthropic-cookbook）— Prompt Engineering、RAG、Agent实战代码')
add_bullet('DeepLearning.AI 的 LangChain短课 — 快速入门Agent和RAG')
add_bullet('Supabase官方教程 — 数据库+向量检索一站式')

add_heading('7.2 必读文档（优先级排序）', level=2)
add_bullet('Next.js App Router文档 — 通读Getting Started + Routing + Data Fetching三个章节')
add_bullet('Claude API文档 / OpenAI API文档 — 重点看Chat Completions、Function Calling、Streaming')
add_bullet('Supabase pgvector文档 — 重点看Vector Search章节')
add_bullet('Vercel AI SDK的useChat + toolInvocations文档')

add_heading('7.3 社区', level=2)
add_bullet('GitHub Discussions / Issues — 遇到问题先搜GitHub，大概率有人遇到过')
add_bullet('Vercel Community / Discord')
add_bullet('知乎/AI相关专栏 — 关注AI应用落地案例而非模型论文')
add_bullet('Twitter/X上关注 @vercel @anthropic @openai — 了解最新动态')

add_heading('7.4 学习期间用好AI', level=2)
add_bullet('把Claude/ChatGPT当你的私人助教。遇到报错贴代码让它分析')
add_bullet('让AI帮你解释看不懂的代码或概念')
add_bullet('让AI帮你Review你的代码，给出改进建议')
add_bullet('但不要直接让它写整个项目 — 你写，它辅助，才能真正学会')

doc.add_page_break()

# ═══════════════════════════════════════════
# 八、求职策略
# ═══════════════════════════════════════════
add_heading('八、求职策略', level=1)

add_heading('8.1 你的核心竞争力', level=2)
add_para('作为应届生，你的优势不在于"精通某个框架"，而在于：')
add_bullet('GitHub上有4个可运行、有README的AI项目')
add_bullet('能独立从零搭一个AI应用并部署上线')
add_bullet('理解AI应用开发的完整链路（Prompt → RAG → Agent → 部署）')
add_bullet('有真实的"踩坑经验"（API限流、Token成本、流式输出、向量检索优化等）')

add_heading('8.2 简历怎么写', level=2)
add_bullet('项目经历放最前面，而非教育背景 — 你的项目比你的学校更值钱')
add_bullet('每个项目写清楚：解决了什么问题、用了什么技术栈、有什么量化成果')
add_bullet('技术栈写具体的：Next.js 14 + Claude API + Supabase pgvector，不要写"熟悉大模型"这种空话')
add_bullet('附GitHub链接和项目Demo链接 — 面试官点开就能用是最好的简历')
add_bullet('如果能写技术博客，链接也附上')

add_heading('8.3 岗位关键词', level=2)
add_para('在招聘平台搜索以下关键词：', font_size=10)
add_bullet('AI应用开发工程师')
add_bullet('大模型应用开发')
add_bullet('全栈工程师（AI方向）')
add_bullet('Prompt Engineer')
add_bullet('RAG开发工程师')
add_bullet('AI产品开发')

add_heading('8.4 面试准备重点', level=2)
add_bullet('AI基础：Token是什么？Temperature怎么影响输出？System Prompt和User Prompt的区别？')
add_bullet('RAG：Embedding原理？Chunking策略？检索质量怎么评估？')
add_bullet('Agent：Function Calling流程？Agent Loop怎么防无限循环？')
add_bullet('工程：流式输出怎么实现？API调用失败怎么重试？Token成本怎么控制？')
add_bullet('系统设计：让你设计一个"企业文档问答系统"，你怎么做？')

add_heading('8.5 投递渠道', level=2)
add_bullet('Boss直聘 / 拉勾 — 国内主力，快速获得面试')
add_bullet('脉脉 — 找人内推，内推成功率远高于海投')
add_bullet('GitHub — 有些公司直接在GitHub Issues招人')
add_bullet('LinkedIn — 外企和远程岗位')
add_bullet('V2EX / 即刻 — 创业公司经常在这招人')

add_heading('8.6 心态提醒', level=2)
add_bullet('不要等"学会了"再投简历 — 第12周就可以开始投，面试本身就是最好的学习')
add_bullet('前几次面试大概率会挂 — 正常，每次面试后记下不会的问题，回去补')
add_bullet('实习比全职容易进 — 先进门，表现好转正')
add_bullet('你不是在和"AI博士"竞争 — 市场上绝大多数AI岗位要的是能落地的人')

add_para('')
add_para('')
add_para('')

# ── 结尾 ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— 现在就开始 —')
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
run.font.size = Pt(14)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('打开终端，输入 npx create-next-app@latest my-ai-app\n注册一个 Claude API 账号，把第一个 AI 聊天跑起来\n遇到问题？问 AI。这就是 AI 应用开发的真实工作方式。')
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# ── 保存 ──
output_dir = os.path.join(os.path.expanduser('~'), 'Desktop')
output_path = os.path.join(output_dir, 'AI应用开发学习计划书.docx')
doc.save(output_path)
import sys
sys.stdout.reconfigure(encoding='utf-8')
print(f'Done: {output_path}')
